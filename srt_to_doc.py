"""
author: Sofia Ojeda
"""

from docx import Document

audio_track = "audio-15"
def srt_to_docx(srt_file, docx_file):
    """
    Converts an SRT subtitle file to a formatted DOCX file.

    :param srt_file: Path to the input SRT file.
    :param docx_file: Path to save the output DOCX file.
    """
    doc = Document()

    with open(f'{audio_track}.srt', "r", encoding="utf-8") as file:
        lines = file.readlines()

    buffer = []
    for line in lines:
        line = line.strip()
        if line.isdigit():
            if buffer:
                timestamp, subtitle_text = process_buffer(buffer)
                write_to_doc(doc, timestamp, subtitle_text)
                buffer = []
        elif line:
            buffer.append(line)
    if buffer:
        timestamp, subtitle_text = process_buffer(buffer)
        write_to_doc(doc, timestamp, subtitle_text)

    doc.save(docx_file)
    print(f"Converted {srt_file} to {docx_file}")

def process_buffer(buffer):
    """
    Extracts the timestamp and subtitle text from a buffer.

    :param buffer: List of lines representing one subtitle entry.
    :return: A tuple of (timestamp, subtitle text).
    """
    timestamp = buffer[0]

    subtitle_text = " ".join(buffer[1:]).replace("<tab>", "\t")
    return timestamp, subtitle_text

def write_to_doc(doc, timestamp, subtitle_text):
    """
    Writes a subtitle entry to the DOCX document.
    Writes timestamp n bold.

    :param doc: The Word Document object.
    :param timestamp: The subtitle timestamp.
    :param subtitle_text: The subtitle text.
    """
    p = doc.add_paragraph()
    p.add_run(timestamp).bold = True
    p.add_run(f"\n{subtitle_text}")

srt_to_docx("audio-15.srt", "audio-15.docx")
